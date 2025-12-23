# # resume_processor.py
# """
# Final resume processing module:
# - Load env
# - Extract contact details (name, email, Indian phone)
# - Extract skills (robust rule-based extractor)
# - Match skills against job-description tokens conservatively
# - Store results via store_files_in_db(cursor, company_name, name, email, phone, skills, file_name, file_path)

# Usage:
#   - Ensure .env is present (sample provided below)
#   - pip install -r requirements (listed below)
#   - import extract_resume_details or run as script for quick self-test
# """

# import os
# import re
# from typing import List, Tuple, Dict
# from difflib import SequenceMatcher

# # dotenv (load .env into os.environ)
# from dotenv import load_dotenv
# load_dotenv()

# # text extraction libraries
# import pdfplumber
# from docx import Document

# # spaCy (NER fallback)
# import spacy

# # DB
# import psycopg2
# from psycopg2 import sql

# # your existing DB storage helper - leave as-is
# from saving import store_files_in_db

# # ---------------------------
# # Load spaCy model (try transformer, fallback to small)
# # ---------------------------
# try:
#     nlp = spacy.load(os.getenv("SPACY_MODEL", "en_core_web_trf"))
# except Exception:
#     try:
#         nlp = spacy.load("en_core_web_sm")
#         print("Warning: TRF model not available — using en_core_web_sm (lighter).")
#     except Exception:
#         nlp = None
#         print("Warning: spaCy NER not available. Regex fallbacks will be used.")

# # ---------------------------
# # TEXT EXTRACTION
# # ---------------------------
# def extract_text_from_pdf(file_path: str) -> str:
#     try:
#         text = ""
#         with pdfplumber.open(file_path) as pdf:
#             for page in pdf.pages:
#                 page_text = page.extract_text()
#                 if page_text:
#                     text += page_text + "\n"
#         return text
#     except Exception as e:
#         print(f"Error extracting text from PDF: {e}")
#         return ""

# def extract_text_from_docx(file_path: str) -> str:
#     try:
#         doc = Document(file_path)
#         text = "\n".join([para.text for para in doc.paragraphs if para.text])
#         return text
#     except Exception as e:
#         print(f"Error extracting text from DOCX: {e}")
#         return ""

# def extract_text_from_file(file_path: str) -> str:
#     if file_path.lower().endswith(".pdf"):
#         return extract_text_from_pdf(file_path)
#     elif file_path.lower().endswith(".docx") or file_path.lower().endswith(".doc"):
#         return extract_text_from_docx(file_path)
#     else:
#         print("Unsupported file format:", file_path)
#         return ""

# # ---------------------------
# # CONTACT EXTRACTION
# # ---------------------------
# def extract_phone_numbers(text: str) -> list:
#     """
#     Robust Indian phone extractor.
#     Normalizes to +91XXXXXXXXXX and returns deduplicated list preserving order.
#     """
#     if not text:
#         return []

#     # match contiguous runs containing digits and common separators
#     pattern = re.compile(r'[\d\+\-\.\(\)\s]{10,30}')
#     candidates = [m.group(0) for m in pattern.finditer(text)]

#     cleaned = []
#     seen = set()
#     for raw in candidates:
#         digits = re.sub(r'\D', '', raw)
#         if len(digits) >= 10:
#             mobile = digits[-10:]
#             # only accept likely Indian mobile numbers that start with 6-9
#             if mobile and mobile[0] in '6789':
#                 normalized = "+91" + mobile
#                 if normalized not in seen:
#                     seen.add(normalized)
#                     cleaned.append(normalized)
#     return cleaned


# def extract_emails(text: str) -> List[str]:
#     emails = re.findall(r"[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}", text)
#     return list(dict.fromkeys([e.lower() for e in emails]))

# def extract_name_by_proximity(text: str, emails: List[str], phones: List[str]) -> str:
#     lines = [ln.strip() for ln in text.splitlines() if ln.strip()]
#     contact_indices = set()
#     for i, ln in enumerate(lines):
#         for e in emails:
#             if e in ln:
#                 contact_indices.add(i)
#         for p in phones:
#             if p.replace("+91", "") in re.sub(r"\D", "", ln):
#                 contact_indices.add(i)
#     for idx in sorted(contact_indices):
#         for j in range(max(0, idx - 3), idx):
#             candidate = lines[j]
#             words = candidate.split()
#             if 1 < len(words) <= 4 and all(w[0].isupper() for w in words if w):
#                 return candidate
#     # spaCy fallback
#     if nlp:
#         doc = nlp(text)
#         persons = [ent.text for ent in doc.ents if ent.label_ == "PERSON"]
#         if persons:
#             return persons[0]
#     return None

# def extract_details(text):
#     emails = extract_emails(text)
#     phones = extract_phone_numbers(text)
#     name = extract_name_by_proximity(text, emails, phones)

#     # Debug prints (remove or turn into logging later)
#     print("DEBUG: extracted emails:", emails)
#     print("DEBUG: extracted phones:", phones)
#     print("DEBUG: extracted name:", name)

#     return name, (emails[0] if emails else None), (phones[0] if phones else None)


# # ---------------------------
# # SKILLS EXTRACTION & MATCHING
# # ---------------------------
# ALIAS_MAP = {
#     "python": {"py", "python", "python3"},
#     "javascript": {"js", "javascript", "ecmascript"},
#     "nodejs": {"node", "nodejs", "node.js"},
#     "java": {"java"},
#     "c++": {"c++", "cpp"},
#     "c#": {"c#", "csharp"},
#     "sql": {"sql"},
#     "postgresql": {"postgresql", "postgres"},
#     "mongodb": {"mongodb", "mongo"},
#     "react": {"react", "reactjs", "react.js"},
#     "docker": {"docker", "docker-compose"},
#     "kubernetes": {"kubernetes", "k8s"},
#     "git": {"git", "github", "gitlab"},
#     "ci/cd": {"ci/cd", "ci", "jenkins"},
#     "aws": {"aws"},
#     "azure": {"azure"},
#     "gcp": {"gcp"},
#     "terraform": {"terraform"},
#     "spring boot": {"spring", "spring boot"},
#     "html": {"html", "html5"},
#     "css": {"css", "css3"},
#     "typescript": {"typescript", "ts"},
# }
# REVERSE_ALIAS: Dict[str, str] = {}
# for canonical, aliases in ALIAS_MAP.items():
#     for a in aliases:
#         REVERSE_ALIAS[a.lower()] = canonical

# NEVER_MATCH_PAIRS = {("java", "javascript"), ("c", "c++")}
# COMMON_JUNK_TOKENS = {"tools", "languages", "familiar", "experienced", "proficiencies", "proficiency", "skills", "skillset"}

# def normalize_token(tok: str) -> str:
#     if tok is None:
#         return ""
#     t = tok.strip().lower()
#     t = re.sub(r"^[^\w\+#\.]+|[^\w\+#\.]+$", "", t)
#     t = re.sub(r"\s+", " ", t)
#     return t

# def canonicalize(tok: str) -> str:
#     t = normalize_token(tok)
#     t = re.sub(r"\.js$", "js", t)
#     t = re.sub(r"\.net$", "dotnet", t)
#     if t in REVERSE_ALIAS:
#         return REVERSE_ALIAS[t]
#     t_stripped = re.sub(r"[^a-z0-9\+#\s\-\.]", "", t)
#     if t_stripped in REVERSE_ALIAS:
#         return REVERSE_ALIAS[t_stripped]
#     return t

# SKILLS_HEADERS = [
#     r"\bskills?\b", r"\btechnical skills\b", r"\bcore competencies\b",
#     r"\bprofessional skills\b", r"\bexpertise\b", r"\btechnical proficiencies\b",
#     r"\btools & technologies\b", r"\bareas of expertise\b", r"\bskillset\b", r"\bskills & tools\b"
# ]
# STOP_HEADERS = [
#     r"\bexperience\b", r"\beducation\b", r"\bprojects\b", r"\bcertifications?\b",
#     r"\bwork history\b", r"\bachievements\b", r"\bsummary\b", r"\bobjective\b",
#     r"\binterests\b", r"\bpublications\b", r"\bpatents\b", r"\bawards\b"
# ]

# def extract_skills_section_text(text: str) -> str:
#     lines = [ln.rstrip() for ln in text.splitlines()]
#     start = -1
#     inline_capture = ""
#     for i, line in enumerate(lines):
#         for h in SKILLS_HEADERS:
#             if re.search(h, line, re.IGNORECASE):
#                 start = i
#                 if ":" in line:
#                     parts = line.split(":", 1)
#                     if parts[1].strip():
#                         inline_capture = parts[1].strip()
#                 break
#         if start != -1:
#             break
#     if start == -1:
#         return ""
#     collected = []
#     if inline_capture:
#         collected.append(inline_capture)
#     for ln in lines[start + 1:]:
#         if any(re.search(sh, ln, re.IGNORECASE) for sh in STOP_HEADERS):
#             break
#         collected.append(ln)
#     return "\n".join(collected).strip()

# def tokenize_skills(section_text: str, split_slash: bool = True) -> List[str]:
#     if not section_text:
#         return []
#     s = re.sub(r"\([^)]*\)", " ", section_text)
#     s = re.sub(r"https?://\S+|www\.\S+", " ", s)
#     raw_tokens = re.split(r"[,\n;/\|•\t]+", s)
#     cleaned = []
#     for t in raw_tokens:
#         if not t:
#             continue
#         t = t.strip()
#         if ":" in t:
#             t = t.split(":", 1)[-1].strip()
#         if not t:
#             continue
#         if split_slash and "/" in t:
#             parts = [p.strip() for p in t.split("/") if p.strip()]
#             cleaned.extend(parts)
#         else:
#             cleaned.append(t)
#     final = []
#     seen = set()
#     for tok in cleaned:
#         tok_norm = tok.strip()
#         if not tok_norm:
#             continue
#         low = tok_norm.lower()
#         if low in COMMON_JUNK_TOKENS:
#             continue
#         if tok_norm not in seen:
#             seen.add(tok_norm)
#             final.append(tok_norm)
#     return final

# def fallback_extract_from_whole_text(text: str) -> List[str]:
#     tokens = re.split(r"[\s,;:/\|\n]+", text)
#     candidates = []
#     seen = set()
#     for t in tokens:
#         if not t or len(t) < 2:
#             continue
#         norm = normalize_token(t)
#         if norm in REVERSE_ALIAS or re.search(r"[\+#]", t) or norm in {"aws", "azure", "gcp", "docker", "kubernetes", "terraform"}:
#             orig = t.strip()
#             if orig not in seen:
#                 seen.add(orig)
#                 candidates.append(orig)
#     return candidates

# def extract_skills_from_text(text: str) -> List[str]:
#     sec = extract_skills_section_text(text)
#     if sec:
#         toks = tokenize_skills(sec, split_slash=True)
#         if toks:
#             return toks
#     return fallback_extract_from_whole_text(text)

# def safe_fuzzy_match(a: str, b: str, min_ratio_short: float = 0.95, min_ratio_long: float = 0.80) -> Tuple[bool, float]:
#     a = normalize_token(a)
#     b = normalize_token(b)
#     ca = canonicalize(a)
#     cb = canonicalize(b)
#     if (ca, cb) in NEVER_MATCH_PAIRS or (cb, ca) in NEVER_MATCH_PAIRS:
#         return False, 0.0
#     ratio = SequenceMatcher(None, a, b).ratio()
#     la, lb = len(a), len(b)
#     min_ratio = min_ratio_long
#     if max(la, lb) <= 6:
#         min_ratio = min_ratio_short
#     if (a in b or b in a) and abs(la - lb) > 3:
#         return False, ratio
#     return (ratio >= min_ratio), ratio

# def match_resume_to_jd(resume_tokens: List[str], jd_tokens: List[str]) -> Dict[str, Tuple[bool, str, str, float]]:
#     jd_canon = {jd: canonicalize(jd) for jd in jd_tokens}
#     resume_canon_map = [(rt, canonicalize(rt)) for rt in resume_tokens]
#     results: Dict[str, Tuple[bool, str, str, float]] = {}
#     for jd, jd_c in jd_canon.items():
#         matched = False
#         match_info = (False, None, None, 0.0)
#         for raw, rc in resume_canon_map:
#             if rc == jd_c:
#                 matched = True
#                 match_info = (True, "exact_canonical", raw, 1.0)
#                 break
#         if matched:
#             results[jd] = match_info
#             continue
#         for raw, rc in resume_canon_map:
#             if normalize_token(raw) == normalize_token(jd):
#                 matched = True
#                 match_info = (True, "exact_normalized", raw, 1.0)
#                 break
#         if matched:
#             results[jd] = match_info
#             continue
#         best = (False, 0.0, None)
#         for raw, rc in resume_canon_map:
#             ok, ratio = safe_fuzzy_match(jd_c, rc)
#             if ok and ratio > best[1]:
#                 best = (True, ratio, raw)
#         if best[0]:
#             match_info = (True, "fuzzy_canonical", best[2], best[1])
#             results[jd] = match_info
#             continue
#         results[jd] = (False, None, None, 0.0)
#     return results

# # ---------------------------
# # MAIN INTEGRATION: process resume and store
# # ---------------------------
# import os
# import psycopg2
# from psycopg2 import sql
# from typing import List, Set, Tuple, Dict

# # Assumes these functions exist in your module (from earlier code)
# # - extract_text_from_file(file_path) -> str
# # - extract_details(text) -> (name, email, phone)
# # - extract_skills_from_text(text) -> List[str]
# # - canonicalize(token) -> str
# # - match_resume_to_jd(resume_tokens, jd_tokens) -> Dict[jd_skill, (matched_bool, method, resume_token, score)]
# # - store_files_in_db(cursor, company_name, name, email, phone, skills_list, file_name, file_path)

# def _get_db_connection():
#     """
#     Prefer DATABASE_URL DSN if present, else use individual DB_* env vars.
#     Returns a psycopg2 connection (caller must close).
#     """
#     db_url = os.getenv("DATABASE_URL")
#     if db_url:
#         return psycopg2.connect(dsn=db_url)
#     # fallback
#     return psycopg2.connect(
#         host=os.getenv("DB_HOST", "localhost"),
#         user=os.getenv("DB_USER", "postgres"),
#         password=os.getenv("DB_PASSWORD", ""),
#         dbname=os.getenv("DB_NAME", "jobs"),
#         port=int(os.getenv("DB_PORT", 5432))
#     )

# def extract_resume_details(file_name: str, file_path: str, company_name: str, require_jd_match: bool = True) -> bool:
#     """
#     Process a single resume file and store it in DB only if it matches at least one JD skill.
#     - file_name, file_path: resume file metadata
#     - company_name: name of DB table (as before)
#     - require_jd_match: when True (default), resume is rejected unless >=1 JD-skill match.
#                         when False, will store resume even if no JD match (keeps previous fallback behavior).
#     Returns True if stored (selected), False otherwise.
#     """
#     # 1. Extract text
#     text = extract_text_from_file(file_path)
#     if not text:
#         print("Text extraction failed.")
#         return False

#     # 2. Extract contact details
#     name, email, phone = extract_details(text)
#     if not name or not email or not phone:
#         print("Name, email, or phone number not found.")
#         return False

#     # 3. Extract resume tokens once (skills section or fallback)
#     resume_tokens = extract_skills_from_text(text)  # e.g. ["JavaScript", "React.js", ...]
#     # If no tokens found and require_jd_match is True, we can early-reject (no skills to match)
#     if not resume_tokens and require_jd_match:
#         print("No skills detected on resume; rejecting (no JD match possible).")
#         return False

#     # 4. DB: open connection and fetch job descriptions for this company
#     conn = None
#     cursor = None
#     try:
#         conn = _get_db_connection()
#         cursor = conn.cursor()

#         # fetch job_title, job_description rows
#         query = sql.SQL("SELECT job_title, job_description FROM {}").format(sql.Identifier(company_name))
#         cursor.execute(query)
#         jobs = cursor.fetchall()  # list of tuples: (job_title, job_description)

#         # 5. Preprocess job descriptions: canonicalize JD skills
#         # Build a list of (job_title, jd_canonical_skills_list)
#         jd_entries: List[Tuple[str, List[str]]] = []
#         for job_title, job_description in jobs:
#             # split on commas (or adjust splitting if your JD format differs)
#             raw_jd_skills = [s.strip() for s in job_description.split(",") if s.strip()]
#             # canonicalize each JD skill
#             jd_canonical = []
#             for s in raw_jd_skills:
#                 c = canonicalize(s)
#                 if c:
#                     jd_canonical.append(c)
#             # dedupe JD canonical list while preserving order
#             seen = set()
#             jd_canonical_unique = []
#             for c in jd_canonical:
#                 if c not in seen:
#                     seen.add(c)
#                     jd_canonical_unique.append(c)
#             if jd_canonical_unique:
#                 jd_entries.append((job_title, jd_canonical_unique))

#         if not jd_entries:
#             # No job descriptions found — depending on policy, either reject or fallback
#             print("No job descriptions found for company; rejecting by default.")
#             return False

#         # 6. For each job, match resume tokens to job's JD skills
#         # We'll choose the job with the largest number of JD-skill matches.
#         best_match_job = None               # (job_title, jd_skills, matched_jd_skill_list)
#         best_match_count = 0

#         # canonicalize resume tokens once for matching where match_resume_to_jd expects raw resume_tokens
#         # (match_resume_to_jd canonicalizes resume tokens internally)
#         for job_title, jd_skills_canonical in jd_entries:
#             # match_resume_to_jd expects jd_tokens list; pass canonical JD tokens
#             matches = match_resume_to_jd(resume_tokens, jd_skills_canonical)
#             matched = [jd for jd, info in matches.items() if info[0]]
#             if matched:
#                 if len(matched) > best_match_count:
#                     best_match_count = len(matched)
#                     best_match_job = (job_title, jd_skills_canonical, matched)

#         # 7. Decide acceptance based on require_jd_match and best_match_job
#         if not best_match_job:
#             # No JD match found
#             if require_jd_match:
#                 print("No JD-matched skills found. Resume Rejected.")
#                 return False
#             else:
#                 # fallback behavior: canonicalize detected tokens and store them
#                 final_skills = []
#                 for t in resume_tokens:
#                     c = canonicalize(t)
#                     if c and c not in final_skills:
#                         final_skills.append(c)
#                 # If still no final_skills, reject
#                 if not final_skills:
#                     print("No detectable skills to store (even after fallback). Resume Rejected.")
#                     return False
#                 # Store fallback skills under a generic 'no_job_match' record (or as per your schema)
#                 print("No JD match, but require_jd_match=False -> storing detected skills.")
#                 store_files_in_db(cursor, company_name, name, email, phone, final_skills, file_name, file_path)
#                 conn.commit()
#                 return True

#         # 8. If we have a best match, store only the matched JD skills (canonical)
#         best_job_title, best_job_jd_skills, matched_jd_skills = best_match_job
#         # dedupe matched_jd_skills and sort/preserve order relative to JD list
#         matched_set = set(matched_jd_skills)
#         ordered_matched = [s for s in best_job_jd_skills if s in matched_set]

#         print(f"Matched {len(ordered_matched)} JD skills for job '{best_job_title}': {ordered_matched}")

#         # 9. Persist to DB
#         store_files_in_db(cursor, company_name, name, email, phone, ordered_matched, file_name, file_path)
#         conn.commit()
#         print(f"File '{file_name}' stored in table '{company_name}' (matched job: {best_job_title}).")
#         return True

#     except Exception as e:
#         # log exception and ensure nothing is left half-committed
#         print(f"Error processing resume details: {e}")
#         if conn:
#             try:
#                 conn.rollback()
#             except Exception:
#                 pass
#         return False
#     finally:
#         if cursor:
#             cursor.close()
#         if conn:
#             conn.close()

# def process_resumes(resume_files: List[Tuple[str, str]], company_name: str):
#     """
#     resume_files: list of (file_name, file_path)
#     """
#     for file_name, file_path in resume_files:
#         print(f"Processing: {file_name}")
#         selected = extract_resume_details(file_name, file_path, company_name, require_jd_match=True)
#         if selected:
#             print(f"Resume Selected: {file_name}")
#         else:
#             print(f"Resume Rejected: {file_name}")

# extract_details.py
# extract_details.py
"""
Final resume processing using skill_matcher.SkillMatcher and skills_config.json.

Usage:
    from extract_details import extract_resume_details, process_resumes
    extract_resume_details("file.pdf", "file.pdf", "MyCompany")  # will store if JD matched
"""

import os
import re
from typing import List, Tuple
import pdfplumber
from docx import Document
import psycopg2
from psycopg2 import sql

# at top of file (near other imports)
try:
    import psycopg2
except ImportError:
    psycopg2 = None

# spaCy (NER fallback)
try:
    import spacy
    nlp = spacy.load(os.getenv("SPACY_MODEL", "en_core_web_sm"))
except Exception:
    nlp = None

from skill_matcher import SkillMatcher
from saving import store_files_in_db  # your existing helper

# load matcher (singleton)
_SKILL_MATCHER = None
def get_matcher():
    global _SKILL_MATCHER
    if _SKILL_MATCHER is None:
        cfg_path = os.getenv("SKILLS_CONFIG", "skills_config.json")
        _SKILL_MATCHER = SkillMatcher(cfg_path)
    return _SKILL_MATCHER

# text extraction helpers
def extract_text_from_pdf(path: str) -> str:
    try:
        text = ""
        with pdfplumber.open(path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
        print (text) 
        return text
    except Exception as e:
        print("PDF text extraction error:", e)
        return ""

def extract_text_from_docx(path: str) -> str:
    try:
        doc = Document(path)
        text = "\n".join([p.text for p in doc.paragraphs if p.text])
        print(text)
        return text
    except Exception as e:
        print("DOCX extraction error:", e)
        return ""

def extract_text_from_file(path: str) -> str:
    p = path.lower()
    if p.endswith(".pdf"):
        return extract_text_from_pdf(path)
    if p.endswith(".docx") or p.endswith(".doc"):
        return extract_text_from_docx(path)
    return ""

# contact extraction
def extract_emails(text: str) -> List[str]:
    if not text:
        return []
    emails = re.findall(r"[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}", text)
    return list(dict.fromkeys([e.lower() for e in emails]))

def extract_phone_numbers(text: str) -> List[str]:
    if not text:
        return []
    # match runs with digits and separators
    candidates = re.findall(r'(\+?\d[\d\-\s\(\)]{8,}\d)', text)
    cleaned = []
    seen = set()
    for raw in candidates:
        digits = re.sub(r'\D', '', raw)
        if len(digits) >= 10:
            mobile = digits[-10:]
            if mobile and mobile[0] in "6789":
                val = "+91" + mobile
                if val not in seen:
                    seen.add(val)
                    cleaned.append(val)
    return cleaned

def extract_name_by_proximity(text: str, emails: List[str], phones: List[str]) -> str:
    lines = [ln.strip() for ln in text.splitlines() if ln.strip()]
    contact_indices = set()
    for i, ln in enumerate(lines):
        for e in emails:
            if e in ln:
                contact_indices.add(i)
        for p in phones:
            if p.replace("+91", "") in re.sub(r'\D', '', ln):
                contact_indices.add(i)
    for idx in sorted(contact_indices):
        for j in range(max(0, idx-3), idx):
            cand = lines[j]
            words = cand.split()
            if 1 < len(words) <= 4 and all(w[0].isupper() for w in words if w):
                return cand
    # spaCy fallback
    if nlp:
        doc = nlp(text)
        persons = [ent.text for ent in doc.ents if ent.label_ == "PERSON"]
        if persons:
            return persons[0]
    return None

# skills section extraction (robust)
SKILLS_HEADERS = [
    r"\bskills?\b", r"\btechnical skills\b", r"\bcore competencies\b",
    r"\bexpertise\b", r"\bproficiencies\b", r"\bskillset\b"
]
STOP_HEADERS = [
    r"\bexperience\b", r"\beducation\b", r"\bprojects\b", r"\bcertifications?\b",
    r"\bwork history\b"
]

def extract_skills_section_text(text: str) -> str:
    lines = text.splitlines()
    start = -1
    inline = ""
    for i, ln in enumerate(lines):
        for h in SKILLS_HEADERS:
            if re.search(h, ln, re.IGNORECASE):
                start = i
                if ":" in ln:
                    parts = ln.split(":", 1)
                    if parts[1].strip():
                        inline = parts[1].strip()
                break
        if start != -1:
            break
    if start == -1:
        return ""
    out = []
    if inline:
        out.append(inline)
    for ln in lines[start+1:]:
        if any(re.search(sh, ln, re.IGNORECASE) for sh in STOP_HEADERS):
            break
        out.append(ln)
    return "\n".join(out).strip()

def tokenize_skills(section_text: str) -> List[str]:
    if not section_text:
        return []
    s = re.sub(r'\([^)]*\)', ' ', section_text)
    s = re.sub(r'https?://\S+|www\.\S+', ' ', s)
    parts = re.split(r'[,\n;/\|\u2022\u2023]+', s)
    tokens = []
    for p in parts:
        p = p.strip()
        if not p:
            continue
        if ":" in p:
            p = p.split(":",1)[-1].strip()
        if "/" in p:
            for x in p.split("/"):
                x = x.strip()
                if x:
                    tokens.append(x)
        else:
            tokens.append(p)
    # dedupe preserve order
    seen = set()
    out = []
    for t in tokens:
        low = t.lower()
        if low in {"skills","skillset","technical"}:
            continue
        if t not in seen:
            seen.add(t)
            out.append(t)
    return out

def fallback_extract_from_whole_text(text: str) -> List[str]:
    tokens = re.split(r'[\s,;:/\|\n]+', text)
    seen = set()
    candidates = []
    for t in tokens:
        if not t or len(t) < 2:
            continue
        if t.isupper() and len(t) <= 4:
            tok = t.strip()
            if tok not in seen:
                seen.add(tok)
                candidates.append(tok)
            continue
        norm = t.strip()
        # heuristic: has + or # or dot (c++, c#, react.js) or contains letters and digits mix (py3)
        if re.search(r'[\+#\.]', norm) or re.search(r'\d', norm):
            if norm not in seen:
                seen.add(norm)
                candidates.append(norm)
            continue
        # if token looks like known skill alias, include
        if norm.lower() in get_matcher().aliases:
            if norm not in seen:
                seen.add(norm)
                candidates.append(norm)
    return candidates

# DB connection helper (uses DATABASE_URL env if present)
def _get_db_connection():
    """
    Prefer DATABASE_URL DSN if present, else use individual DB_* env vars.
    Returns a psycopg2 connection (caller must close).
    Raises RuntimeError with clear guidance if psycopg2 is not installed.
    """
    if psycopg2 is None:
        raise RuntimeError("psycopg2 is not installed. Install it with: pip install psycopg2-binary")

    db_url = os.getenv("DATABASE_URL")
    if db_url:
        return psycopg2.connect(dsn=db_url)

    # fallback to env vars
    return psycopg2.connect(
        host=os.getenv("DB_HOST", "localhost"),
        user=os.getenv("DB_USER", "postgres"),
        password=os.getenv("DB_PASSWORD", "Ayush123"),
        dbname=os.getenv("DB_NAME", "jobs"),
        port=int(os.getenv("DB_PORT", 5432))
    )

def extract_resume_details(file_name: str, file_path: str, company_name: str, require_jd_match: bool = True) -> bool:
    """
    Extract and store resume only if matches JD skills stored in DB table company_name.
    Returns True if stored, False otherwise.
    """
    text = extract_text_from_file(file_path)
    if not text:
        print("Text extraction failed.")
        return False

    emails = extract_emails(text)
    phones = extract_phone_numbers(text)
    name = extract_name_by_proximity(text, emails, phones)

    print("DEBUG: extracted emails:", emails)
    print("DEBUG: extracted phones:", phones)
    print("DEBUG: extracted name:", name)

    if not name or not emails or not phones:
        print("Name, email, or phone missing. Rejecting.")
        return False

    # Extract skill tokens
    sec = extract_skills_section_text(text)
    if sec:
        resume_tokens = tokenize_skills(sec)
        if not resume_tokens:
            resume_tokens = fallback_extract_from_whole_text(text)
    else:
        resume_tokens = fallback_extract_from_whole_text(text)

    if not resume_tokens:
        print("No skill tokens detected.")
        if require_jd_match:
            return False

    matcher = get_matcher()
    # Connect DB and fetch JDs
    conn = None
    cursor = None
    try:
        conn = _get_db_connection()
        cursor = conn.cursor()
        query = sql.SQL("SELECT job_title, job_description FROM {}").format(sql.Identifier(company_name))
        cursor.execute(query)
        jobs = cursor.fetchall()
        if not jobs:
            print("No jobs found for company:", company_name)
            return False

        best_job = None
        best_matches = {}
        best_count = 0

        for job_title, job_description in jobs:
            # JD parse: split on comma
            jd_raw = [s.strip() for s in job_description.split(",") if s.strip()]
            matches = matcher.match_resume_to_jd(resume_tokens, jd_raw)
            matched = [jd for jd,info in matches.items() if info[0]]
            print(f"Matched {len(matched)} JD skills for job '{job_title}': {matched}")
            if len(matched) > best_count:
                best_count = len(matched)
                best_job = (job_title, jd_raw, matches)
                best_matches = matches

        if not best_job:
            print("No JD-matched skills found. Resume Rejected.")
            return False

        # Prepare matched canonical list for storing: use JD items that matched (canonicalized)
        matched_jd_skills = [jd for jd, info in best_matches.items() if info[0]]
        if not matched_jd_skills:
            print("No matched JD skills after evaluation.")
            return False

        print(f"Matched {len(matched_jd_skills)} JD skills for job '{best_job[0]}': {matched_jd_skills}")

        # Store matched_jd_skills in DB using your helper
        store_files_in_db(cursor, company_name, name, emails[0], phones[0], matched_jd_skills, file_name, file_path)
        conn.commit()
        print(f"File '{file_name}' stored in table '{company_name}' (matched job: {best_job[0]}).")
        return True

    except Exception as e:
        print("Error processing resume:", e)
        if conn:
            try:
                conn.rollback()
            except Exception:
                pass
        return False
    finally:
        if cursor:
            cursor.close()
        if conn:
            conn.close()

def process_resumes(resume_files: List[Tuple[str,str]], company_name: str):
    for file_name, file_path in resume_files:
        print("Processing:", file_name)
        ok = extract_resume_details(file_name, file_path, company_name, require_jd_match=True)
        print("Selected" if ok else "Rejected:", file_name)
